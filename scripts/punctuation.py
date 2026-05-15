#!/usr/bin/env python3
"""
标点符号修复 - GB/T 9704-2012 公文标准
来源: https://github.com/cj0103/gbt-9704-2012-skills

用法: python punctuation.py input.docx output.docx [--mode smart|chinese] [--fix brackets,quotes]
"""

import re
from docx import Document

LEFT_DOUBLE_QUOTE = '\u201c'
RIGHT_DOUBLE_QUOTE = '\u201d'
LEFT_SINGLE_QUOTE = '\u2018'
RIGHT_SINGLE_QUOTE = '\u2019'

REPLACEMENTS = {
    "(": "（", ")": "）", ":": "：", ";": "；", "?": "？", "!": "！",
}

def has_chinese(text):
    return bool(re.search(r"[\u4e00-\u9fff]", text))

def fix_text(text, mode='smart'):
    if not text:
        return text
    result = text
    if mode == 'chinese':
        for en, cn in REPLACEMENTS.items():
            result = result.replace(en, cn)
        result = re.sub(r",", "，", result)
        result = re.sub(r"\.(?!\.)", "。", result)
    elif mode == 'english':
        for cn, en in [("（", "("), ("）", ")"), ("：", ":"), ("；", ";"), ("，", ","), ("。", ".")]:
            result = result.replace(cn, en)
        result = result.replace("……", "...").replace("——", "--")
    else:
        result = re.sub(r"\.{2,}", "……", result)
        result = re.sub(r"。{2,}", "……", result)
        result = re.sub(r"--+", "——", result)
        result = re.sub(r"—(?!—)", "——", result)
        if has_chinese(result):
            for en, cn in REPLACEMENTS.items():
                result = result.replace(en, cn)
        result = re.sub(r"([\u4e00-\u9fff]),", r"\1，", result)
        result = re.sub(r",([\u4e00-\u9fff])", r"，\1", result)
        result = re.sub(r"([\u4e00-\u9fff])\.(\s|$)", r"\1。\2", result)
    double_chars = ['"', '\u201c', '\u201d', '\u201e', '\u201f', '\u300c', '\u300d']
    temp_result = result
    for q in double_chars:
        temp_result = temp_result.replace(q, "\x00")
    if "\x00" in temp_result:
        chars = list(temp_result)
        quote_count = 0
        for i, c in enumerate(chars):
            if c == "\x00":
                if quote_count % 2 == 0:
                    chars[i] = LEFT_DOUBLE_QUOTE
                else:
                    chars[i] = RIGHT_DOUBLE_QUOTE
                quote_count += 1
        result = "".join(chars)
    single_chars = ["'", '\u2018', '\u2019', '\u201a', '\u201b']
    temp_result = result
    for q in single_chars:
        temp_result = temp_result.replace(q, "\x01")
    if "\x01" in temp_result:
        chars = list(temp_result)
        quote_count = 0
        for i, c in enumerate(chars):
            if c == "\x01":
                if quote_count % 2 == 0:
                    chars[i] = LEFT_SINGLE_QUOTE
                else:
                    chars[i] = RIGHT_SINGLE_QUOTE
                quote_count += 1
        result = "".join(chars)
    return result

def process_paragraph(para, mode='smart'):
    full_text = para.text
    if not full_text.strip():
        return False
    fixed_text = fix_text(full_text, mode)
    if fixed_text == full_text:
        return False
    runs = para.runs
    if not runs:
        return False
    first_run = runs[0]
    first_run.text = fixed_text
    for run in runs[1:]:
        run.text = ""
    return True

def process_document(input_path, output_path, mode='smart', fix_types=None):
    print(f"Reading: {input_path}")
    doc = Document(input_path)
    changes = 0
    for i, para in enumerate(doc.paragraphs):
        if process_paragraph(para, mode):
            changes += 1
            preview = para.text[:50] + "..." if len(para.text) > 50 else para.text
            print(f"  Para {i + 1}: {preview}")
    table_changes = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if process_paragraph(para, mode):
                        table_changes += 1
    if table_changes:
        print(f"  Tables: {table_changes} cells fixed")
    print(f"\nTotal: {changes} paragraphs + {table_changes} table cells fixed")
    doc.save(output_path)
    print(f"Saved: {output_path}")

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description='公文标点符号修复工具')
    parser.add_argument('input', help='输入文件路径')
    parser.add_argument('output', help='输出文件路径')
    parser.add_argument('--mode', choices=['smart', 'chinese', 'english'], default='smart')
    parser.add_argument('--fix', help='只修复特定类型，如：brackets,quotes')
    args = parser.parse_args()
    process_document(args.input, args.output, mode=args.mode)
